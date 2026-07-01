import io

import docx
import pdfplumber


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a .docx file."""
    document = docx.Document(io.BytesIO(file_bytes))
    
    parts: list[str] = []
    
    # Extract from paragraphs
    for p in document.paragraphs:
        text = p.text.strip()
        if text:
            parts.append(text)
    
    # Extract from tables
    for table in document.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))
    
    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a .pdf file."""
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
    return "\n".join(text_parts)
